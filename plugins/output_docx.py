# -*- coding: utf-8 -*-
"""
@Time ： 2022/3/15 5:28 下午
@Author ： Shawn
@FileName ：output_docx.py
@IDE ：PyCharm

"""

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import os
from _datetime import datetime

path = os.path.abspath(__file__)
os = os.path.split(path)[0]
SerInfo = ['PHPSESSION','ASP.NET_SessionId','JSESSIONI']
##创建文档
document = Document()
#添加标题
def AddHeadText(text, size):
    title_ = document.add_heading(level=0)
    title_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER# 标题居中
    title_run = title_.add_run(text)  # 添加标题内容
    title_run.font.size = Pt(size)  # 设置标题字体大小
    title_run.font.name = 'Times New Roman'
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体') # 设置标题黑体字体
    title_run.font.color.rgb = RGBColor(0, 0, 0)#字体颜色

## 添加段落，(参数1：文本内容，参数2：字体大小，参数3：上行距,参数4：字体粗细，参数5：段落位置)
def AddParaText(text, size, space, thickness, position):
    p = document.add_paragraph()  # 段落
    #判断居中还是靠左,0为靠左
    if position == 0:
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT #靠左
    else :
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
    p.paragraph_format.space_before = Pt(space)
    text = p.add_run(text)
    #判断字体是否加粗（1为不加粗）
    if thickness == 1:
        text.bold = False
    else:
        text.bold = True #加粗
    text.font.name = 'Times New Roman'
    text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    text.font.size = Pt(size)

def create_Serinfo_table(rows,columns,result):
    AddParaText("        1、详情", 13, 5, 0, 0)
    # rows 行    cols    列
    tables = document.add_table(rows=rows,cols=columns,style="Light Grid Accent 5")
    url = []
    res = []
    ServerInfo = ''
    url_info = 0
    urls = 0
    title = [
        ['编号','URL地址','详情','安全分析']
    ]
    for u in result:
        url.append(u)
        res.append(result[u])
    ## 写表头,行
    for i in range(rows):
        cells = tables.rows[i].cells
        ## 写表内容，列
        for r in range(columns):
            if i == 0:
                cells[r].text = str(title[i][r])
            # 编号
            elif i > 0 and r == 0:
                cells[r].text = str(i)
            # url
            elif i > 0 and r == 1:
                cells[r].text = url[urls]
                urls += 1
            # ServerInfo
            elif i > 0 and r == 2:
                cells[r].text = res[url_info]
                if res[url_info] in SerInfo:
                    ServerInfo = "不安全，建议对服务器配置进行修改"
                elif res[url_info] == 'None':
                    ServerInfo = "未匹配到字段，可能情况为:\n1、服务器配置安全性高。\n2、Url地址错误。\n3、服务器已失活不可访问。\n4、当前代理IP已被封禁"
                else:
                    ServerInfo = '安全，暂未发现服务器信息泄漏'
                url_info += 1
            # 信息分析
            elif i > 0 and r == 3:
                cells[r].text = ServerInfo
    AddParaText("        2、简要分析", 13, 5, 0, 0)
    AddParaText("        信息泄露：暴露服务器的信息，使攻击者能够通过泄露的信息入侵。",12,5,5,0)

def create_CmsInfo_table(rows,columns,result):
    AddParaText("        1、详情", 13, 5, 0, 0)
    url = []
    res = []
    # rows 行    cols    列
    tables = document.add_table(rows=rows, cols=columns, style="Light Grid Accent 5")
    title = [
        ['编号', 'URL地址', 'CMS识别']
    ]
    CMS_info = 0
    urls = 0
    for u in result:
        ## 获取键值
        url.append(u)
        res.append(result[u])
    ## 写表头,行
    for i in range(rows):
        cells = tables.rows[i].cells
        ## 写表内容，列
        for r in range(columns):
            if i == 0:
                cells[r].text = str(title[i][r])
            # 编号
            elif i > 0 and r == 0:
                cells[r].text = str(i)
            # url
            elif i > 0 and r == 1:
                cells[r].text = url[urls]
                urls += 1
            # CMS_INFO
            elif i > 0 and r == 2:
                cells[r].text = res[CMS_info]
                CMS_info += 1
    AddParaText("        2、简要分析", 13, 5, 0, 0)
    AddParaText("        CMS系统信息：由于大多数的CMS系统为开源系统，其代码均公开于互联网。防护建议：建议更新至官方最新版本。", 12, 5, 5, 0)

def check_content(result,type):
    if type == 'ServerInfo':
        output_addr = ServerInfo_add(result)
        return output_addr
    if type == 'CMS_info':
        return CMS_Search(result)
    if type == 'WhoisInfo':
        return WhoisInfo(result)
    if type == 'Subdomain':
        return Subdomain(result)


def save_file():
    document.save(os + '/../result/' + datetime.now().strftime("%Y%m%d_%H%M%S") + '.docx')
    return  os + '/../result/' + datetime.now().strftime("%Y%m%d_%H%M%S") + '.docx'


def ServerInfo_add(result):
    columns = 4
    AddParaText("服务器信息收集", 15, 5, 0, 0)
    rows = len(result) + 1
    create_Serinfo_table(rows,columns,result)
    return save_file()
def CMS_Search(result):
    columns = 3
    AddParaText("指纹识别信息收集", 15, 5, 0, 0)
    rows = len(result) + 1
    create_CmsInfo_table(rows,columns,result)
    return save_file()
def WhoisInfo(result):
    def finish_whois():
        AddParaText("        2、简要分析", 13, 5, 0, 0)
        AddParaText("        CMS系统信息：由于大多数的CMS系统为开源系统，其代码均公开于互联网。防护建议：建议更新至官方最新版本。", 12, 5, 5, 0)

    def create_whois_table(rows,columns,result,addr):
        wh = []
        icp = []
        whois = ''
        icps = ''
        tables = document.add_table(rows=rows, cols=columns, style="Light Grid Accent 5")
        # 遍历whois字典键值,除掉end,url,addr字段
        for w in result[0]:
            if w == 'end':
                pass
            elif w == 'url':
                pass
            elif w == 'addr':
                pass
            else:
                wh.append(w)
        for a in range(len(wh)):
            if a == 0:
                whois = str(wh[a]) + ':' + str(result[0][wh[a]]) + '\n'
            elif wh[a] == '域名状态':
                c = ''
                for aa in range(len(result[0][wh[a]])):
                    if aa == 0:
                        c = str(result[0][wh[a]][aa])
                    else:
                        c = c + str(result[0][wh[a]][aa])
                whois = str(whois) + str(wh[a]) + ':' + c + '\n'
            else:
                whois = str(whois) + str(wh[a]) + ':' + str(result[0][wh[a]]) + '\n'
        # 遍历icp字典键值
        if result[1] != 0:
            for ic in result[1]:
                icp.append(ic)
            for b in range(len(icp)):
                if b == 0:
                    icps = str(icp[b]) + ':' + str(result[1][icp[b]]) + '\n'
                else:
                    icps = str(icps) + str(icp[b]) + ':' + str(result[1][icp[b]]) + '\n'
        else:
            icps = '暂无'
        for i in range(rows):
            cells = tables.rows[i].cells
            ## 写表内容，列
            for r in range(columns):
                if i == 0 and result[0]['end'] == 0:
                    cells[r].text = str(title[i][r])
                elif i > 0 and r == 0:
                    cells[r].text = result[0]['url']
                # url
                elif i > 0 and r == 1:
                    cells[r].text = whois
                if i > 0 and r == 2:
                    cells[r].text = icps
        try:
            if result[0]['end'] == 0:
                document.save(os + '/../result/' + datetime.now().strftime("%Y%m%d_%H%M%S") + '.docx')
                save_addr = os + '/../result/' + datetime.now().strftime("%Y%m%d_%H%M%S") + '.docx'
                return save_addr
            elif result[0]['end'] == 'end':
                document.save(addr)
                finish_whois()
                return addr
            else:
                document.save(addr)
        except Exception:
            return 'Save File Error'

    columns = 3
    rows = 2
    title = [
        ['Url地址', 'Whois信息', 'ICP备案信息']
    ]
    addr = result[0]['addr']
    if result[0]['end'] == 0:
        AddParaText("Whois/备案信息收集", 15, 5, 0, 0)
    return create_whois_table(rows,columns,result,addr)

def Subdomain(result):
    AddParaText("子域名信息收集", 15, 5, 0, 0)
    url = []
    res = ''
    # 遍历键值
    for o in result:
        url.append(o)
        for u in range(len(result[o])):
            if u == 0:
                res = result[o][0]
            else:
                res = res + '\n' + result[o][u]
    AddParaText("        1、详情", 13, 5, 0, 0)
    # rows 行    cols    列
    tables = document.add_table(rows=len(url) + 1, cols=3, style="Light Grid Accent 5")
    title = [
        ['编号', 'URL地址', '子域名']
    ]
    for i in range(len(url)+1):
        cells = tables.rows[i].cells
        ## 写表内容，列
        for r in range(3):
            if i == 0:
                cells[r].text = str(title[i][r])
            elif i > 0 and r == 0:
                cells[r].text = str(i)
            # url
            elif i > 0 and r == 1:
                if i < len(url)+1:
                    cells[r].text = url[i-1]
            if i > 0 and r == 2:
                cells[r].text = res
    AddParaText("        子域名资产的暴露会增大被攻击面，从而使得攻击成功率的提升，防御方式：添加泛域名，禁止部分IP访问，对子域名站点购买防御设备，增加子域名的命名复杂度。", 12, 5, 5, 0)
    return save_file()

AddHeadText('信息收集测试报告',20)
# ServerInfo_add()
